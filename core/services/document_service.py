"""Document Parsing Service
Extracts text from various document formats: TXT, MD, DOCX, PDF
"""

import logging
import os
from typing import Optional, Tuple
from pathlib import Path

import chardet
from pypdf import PdfReader
from docx import Document
import markdown2

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse text from various document formats"""
    
    SUPPORTED_EXTENSIONS = ['.txt', '.md', '.docx', '.pdf', '.markdown']
    
    @staticmethod
    def is_supported(file_path: str) -> bool:
        """Check if file format is supported"""
        ext = Path(file_path).suffix.lower()
        return ext in DocumentParser.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def detect_encoding(file_path: str) -> str:
        """Detect file encoding"""
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)  # Read first 10KB
                result = chardet.detect(raw_data)
                return result['encoding'] or 'utf-8'
        except Exception as e:
            logger.warning(f"Failed to detect encoding for {file_path}: {e}")
            return 'utf-8'
    
    @staticmethod
    def parse_txt(file_path: str) -> str:
        """Parse plain text file"""
        try:
            encoding = DocumentParser.detect_encoding(file_path)
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                text = f.read()
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing TXT file {file_path}: {e}")
            raise
    
    @staticmethod
    def parse_markdown(file_path: str) -> str:
        """Parse markdown file (returns plain text without HTML)"""
        try:
            encoding = DocumentParser.detect_encoding(file_path)
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                md_content = f.read()
            
            # Convert markdown to HTML then strip HTML tags to get plain text
            html = markdown2.markdown(md_content)
            
            # Simple HTML tag removal
            import re
            text = re.sub(r'<[^>]+>', '', html)
            text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing Markdown file {file_path}: {e}")
            raise
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Parse DOCX file"""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        tables_text.append(' | '.join(row_text))
            
            # Combine all text
            all_text = paragraphs + tables_text
            text = '\n\n'.join(all_text)
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {e}")
            raise
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Parse PDF file"""
        try:
            reader = PdfReader(file_path)
            
            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")
                    continue
            
            if not text_parts:
                raise ValueError("No text could be extracted from PDF")
            
            text = '\n\n'.join(text_parts)
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing PDF file {file_path}: {e}")
            raise
    
    @staticmethod
    def parse_document(file_path: str) -> Tuple[str, str]:
        """
        Parse document and return (text, format)
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (extracted_text, file_format)
            
        Raises:
            ValueError: If file format is not supported
            Exception: If parsing fails
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = Path(file_path).suffix.lower()
        
        if ext not in DocumentParser.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                f"Supported formats: {', '.join(DocumentParser.SUPPORTED_EXTENSIONS)}"
            )
        
        logger.info(f"Parsing document: {file_path} (format: {ext})")
        
        # Parse based on extension
        if ext == '.txt':
            text = DocumentParser.parse_txt(file_path)
            format_name = "TXT"
        elif ext in ['.md', '.markdown']:
            text = DocumentParser.parse_markdown(file_path)
            format_name = "Markdown"
        elif ext == '.docx':
            text = DocumentParser.parse_docx(file_path)
            format_name = "DOCX"
        elif ext == '.pdf':
            text = DocumentParser.parse_pdf(file_path)
            format_name = "PDF"
        else:
            raise ValueError(f"Unsupported extension: {ext}")
        
        if not text or len(text.strip()) < 50:
            raise ValueError(f"Document appears to be empty or too short (< 50 characters)")
        
        logger.info(f"Successfully parsed {format_name} document: {len(text)} characters")
        
        return text, format_name


def extract_text_from_document(file_path: str) -> Optional[str]:
    """
    Convenience function to extract text from document
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Extracted text or None if parsing fails
    """
    try:
        text, _ = DocumentParser.parse_document(file_path)
        return text
    except Exception as e:
        logger.error(f"Failed to extract text from {file_path}: {e}")
        return None


def get_document_info(file_path: str) -> dict:
    """
    Get document information without parsing full content
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Dictionary with document metadata
    """
    if not os.path.exists(file_path):
        return {"error": "File not found"}
    
    path_obj = Path(file_path)
    stat = os.stat(file_path)
    
    return {
        "filename": path_obj.name,
        "extension": path_obj.suffix.lower(),
        "size_bytes": stat.st_size,
        "size_kb": round(stat.st_size / 1024, 2),
        "is_supported": DocumentParser.is_supported(file_path)
    }
